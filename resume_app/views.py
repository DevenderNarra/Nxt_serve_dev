import boto3
import json
from io import BytesIO
from docx import Document
from pdfminer.high_level import extract_text
from django.conf import settings
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import Position, Candidate
import google.generativeai as genai
from django.shortcuts import render
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from rest_framework.permissions import AllowAny
from rest_framework import status
from rest_framework_simplejwt.views import TokenObtainPairView
from .serializers import EmployerSignupSerializer, InterviewerSignupSerializer, MyTokenObtainPairSerializer

class LoginView(TokenObtainPairView):
    permission_classes = [AllowAny]
    serializer_class = MyTokenObtainPairSerializer

class EmployerSignupView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = EmployerSignupSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Employer registered successfully"}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class InterviewerSignupView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = InterviewerSignupSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Interviewer registered successfully"}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

def test_position(request):
    return render(request, 'resume_app/test_position.html')

def test_candidate(request):
    return render(request, 'resume_app/test_candidate.html')

def test_employer_signup(request):
    return render(request, 'resume_app/employer_signup.html')

def test_interviwer_signup(request):
    return render(request, 'resume_app/interviewer_signup.html')

def test_jd(request):
    return render(request, 'resume_app/jd.html')

# ------------ S3 Upload Function ------------
def upload_to_s3(file_obj, file_name):
    s3 = boto3.client(
        's3',
        aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
        aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
        region_name=settings.AWS_S3_REGION_NAME
    )
    s3.upload_fileobj(file_obj, settings.AWS_STORAGE_BUCKET_NAME, file_name)
    url = f"https://{settings.AWS_STORAGE_BUCKET_NAME}.s3.{settings.AWS_S3_REGION_NAME}.amazonaws.com/{file_name}"
    return url


# ------------ Extract Text from File ------------
def extract_text_from_file(file_obj, file_name):
    text = ""
    if file_name.endswith('.pdf'):
        text = extract_text(file_obj)
    elif file_name.endswith('.docx'):
        doc = Document(file_obj)
        text = "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")
    return text


# ------------ LLM Resume Parsing ------------
def extract_resume_data_with_llm(resume_text):
    genai.configure(api_key=settings.GEMINI_API_KEY)
    model = genai.GenerativeModel("models/gemini-2.5-flash")

    prompt = f"""
    You are a resume parsing AI. From the given resume text, extract structured information.
    Return JSON with these exact keys:
    - name
    - email
    - phone
    - skills (list)
    - experience

    Resume text:
    {resume_text}
    """

    response = model.generate_content(prompt)
    result_text = response.text.strip()

    try:
        data = json.loads(result_text)
    except json.JSONDecodeError:
        try:
            json_start = result_text.find('{')
            json_end = result_text.rfind('}') + 1
            data = json.loads(result_text[json_start:json_end])
        except Exception:
            data = {"name": "", "email": "", "phone": "", "skills": [], "experience": ""}
    return data


# ------------ LLM JD Parsing ------------
def extract_skills_from_jd(jd_text):
    genai.configure(api_key=settings.GEMINI_API_KEY)
    model = genai.GenerativeModel("models/gemini-2.5-flash")

    prompt = f"""
    You are a job description parsing AI. From the given job description text, extract structured information about skills required for the role.

    Return JSON with these exact keys:
    - mandatory_skills (list of skills that are required)
    - optional_skills (list of skills that are nice-to-have)

    Job Description text:
    {jd_text}
    """

    response = model.generate_content(prompt)
    result_text = response.text.strip()

    try:
        data = json.loads(result_text)
    except json.JSONDecodeError:
        try:
            json_start = result_text.find('{')
            json_end = result_text.rfind('}') + 1
            data = json.loads(result_text[json_start:json_end])
        except Exception:
            data = {"mandatory_skills": [], "optional_skills": []}
    return data


# ------------ API: Create Position ------------
class PositionCreateView(APIView):
    def post(self, request):
        try:
            file = request.FILES.get('job_description_file')
            if not file:
                return Response({"error": "JD file is required."}, status=status.HTTP_400_BAD_REQUEST)

            # Upload JD to S3
            s3_url = upload_to_s3(BytesIO(file.read()), file.name)

            # Extract JD text
            file.seek(0)
            jd_text = extract_text_from_file(BytesIO(file.read()), file.name)

            # Extract mandatory & optional skills using LLM
            skills_data = extract_skills_from_jd(jd_text)
            mandatory_skills = skills_data.get('mandatory_skills', [])
            optional_skills = skills_data.get('optional_skills', [])

            # Create Position
            position = Position.objects.create(
                job_title=request.data.get('job_title'),
                domain=request.data.get('domain'),
                exp_from=request.data.get('exp_from'),
                exp_to=request.data.get('exp_to'),
                job_description_file=file,
                mandatory_skills=mandatory_skills,
                optional_skills=optional_skills,
                interview_instructions=request.data.get('interview_instructions')
            )

            return Response({
                "id": position.id,
                "job_title": position.job_title,
                "domain": position.domain,
                "exp_from": position.exp_from,
                "exp_to": position.exp_to,
                "mandatory_skills": position.mandatory_skills,
                "optional_skills": position.optional_skills,
                "jd_file_url": s3_url,
                "interview_instructions": position.interview_instructions
            }, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    def get(self, request):
        positions = Position.objects.all()
        data = [
            {
                "id": pos.id,
                "job_title": pos.job_title,
                "domain": pos.domain,
                "exp_from": pos.exp_from,
                "exp_to": pos.exp_to,
                "mandatory_skills": pos.mandatory_skills,
                "optional_skills": pos.optional_skills,
                "jd_file_url": pos.job_description_file.url if pos.job_description_file else "",
                "interview_instructions": pos.interview_instructions,
            }
            for pos in positions
        ]
        return Response(data, status=status.HTTP_200_OK)


# ------------ API: Get Position Details ------------
class PositionDetailView(APIView):
    def get(self, request, pk):
        try:
            position = Position.objects.get(id=pk)
            return Response({
                "id": position.id,
                "job_title": position.job_title,
                "domain": position.domain,
                "mandatory_skills": position.mandatory_skills,
                "optional_skills": position.optional_skills,
                "jd_file_url": position.job_description_file.url,
                "interview_instructions": position.interview_instructions
            })
        except Position.DoesNotExist:
            return Response({"error": "Position not found."}, status=status.HTTP_404_NOT_FOUND)


# ------------ API: Add Candidate ------------
class CandidateCreateView(APIView):
    def post(self, request):
        try:
            position_id = request.data.get('position_id')
            if not position_id:
                return Response({"error": "Position ID is required."}, status=status.HTTP_400_BAD_REQUEST)

            try:
                position = Position.objects.get(id=position_id)
            except Position.DoesNotExist:
                return Response({"error": "Position not found."}, status=status.HTTP_404_NOT_FOUND)

            resume_file = request.FILES.get('resume_file')
            if not resume_file:
                return Response({"error": "Resume file is required."}, status=status.HTTP_400_BAD_REQUEST)

            # Upload resume to S3
            s3_url = upload_to_s3(BytesIO(resume_file.read()), resume_file.name)

            # Reset pointer and extract text
            resume_file.seek(0)
            resume_text = extract_text_from_file(BytesIO(resume_file.read()), resume_file.name)

            # Extract structured data from resume using LLM
            extracted_data = extract_resume_data_with_llm(resume_text)

            # Safely get experience as float
            experience_value = request.data.get('experience') or extracted_data.get('experience', 0)
            try:
                experience_value = float(experience_value)
            except (ValueError, TypeError):
                experience_value = 0.0

            # Create candidate
            candidate = Candidate.objects.create(
                position=position,
                domain=request.data.get('domain', position.domain),
                mandatory_skills=request.data.get('mandatory_skills', position.mandatory_skills),
                optional_skills=request.data.get('optional_skills', position.optional_skills),
                jd_file=position.job_description_file,
                resume_file=resume_file,
                name=extracted_data.get('name', request.data.get('name', '')),
                email=extracted_data.get('email', request.data.get('email', '')),
                contact=extracted_data.get('phone', request.data.get('contact', '')),
                experience=experience_value,
                preferred_timings=request.data.get('preferred_timings'),
                interview_instructions=position.interview_instructions
            )

            return Response({
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "contact": candidate.contact,
                "experience": candidate.experience,
                "mandatory_skills": candidate.mandatory_skills,
                "optional_skills": candidate.optional_skills,
                "resume_s3_url": s3_url
            }, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
    # Generate JD
# Configure your GenAI API key
genai.configure(api_key=settings.GEMINI_API_KEY)
model = genai.GenerativeModel("models/gemini-2.5-flash")


# @method_decorator(csrf_exempt, name='dispatch')
# class GenerateJDPreviewView(APIView):
#     def post(self, request):
#         try:
#             job_title = request.data.get('job_title')
#             domain = request.data.get('domain')
#             experience = request.data.get('experience')

#             if not job_title or not domain or experience is None:
#                 return Response(
#                     {"error": "job_title, domain and experience are required."},
#                     status=status.HTTP_400_BAD_REQUEST
#                 )

#             prompt = f"Write a professional job description for a {job_title} in {domain} with {experience} years of experience."
#             response = model.generate_content(prompt=prompt, max_output_tokens=600)
#             jd_text = getattr(response, 'output_text', '')
#             if not jd_text:
#                 return Response({"error": "Failed to generate JD"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#             return Response({"jd_text": jd_text})

#         except Exception as e:
#             return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class GenerateJDFileView(APIView):
    """
    Generate JD as PDF or DOCX file for download.
    """
    def post(self, request):
        try:
            job_title = request.data.get('job_title')
            domain = request.data.get('domain')
            experience = request.data.get('experience')
            output_format = request.data.get('format', 'pdf')  # pdf or docx

            if not job_title or not domain or experience is None:
                return Response(
                    {"error": "job_title, domain and experience are required."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            prompt = f"Write a professional job description for a {job_title} in {domain} with {experience} years of experience."
            response = model.generate_content(prompt=prompt, max_output_tokens=600)
            jd_text = response.output_text

            buffer = BytesIO()

            if output_format.lower() == 'pdf':
                p = canvas.Canvas(buffer)
                y = 800
                for line in jd_text.split('\n'):
                    p.drawString(50, y, line)
                    y -= 15
                p.showPage()
                p.save()
                buffer.seek(0)
                response_file = HttpResponse(buffer, content_type='application/pdf')
                response_file['Content-Disposition'] = f'attachment; filename="{job_title}_JD.pdf"'
            else:  # DOCX
                doc = Document()
                doc.add_heading(f"{job_title} - Job Description", 0)
                doc.add_paragraph(jd_text)
                doc.save(buffer)
                buffer.seek(0)
                response_file = HttpResponse(
                    buffer,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response_file['Content-Disposition'] = f'attachment; filename="{job_title}_JD.docx"'

            return response_file

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
class GenerateJDPreviewView(APIView):
    def post(self, request):
        try:
            job_title = request.data.get('job_title', '').strip()
            domain = request.data.get('domain', '').strip()
            experience = request.data.get('experience', '')

            if not job_title or not domain or not experience:
                return Response(
                    {"error": "job_title, domain and experience are required."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Ensure experience is integer
            try:
                experience = int(experience)
            except ValueError:
                return Response({"error": "experience must be a number."}, status=status.HTTP_400_BAD_REQUEST)

            prompt_text = f"Write a professional job description for a {job_title} in {domain} with {experience} years of experience."

            # ✅ Correct usage for your SDK: only pass the prompt
            response = model.generate_content(prompt_text)

            # Safely get output text
            jd_text = getattr(response, 'output_text', None) or getattr(response, 'text', None)
            if not jd_text:
                return Response({"error": "Failed to generate JD from model."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Send JD as list of lines for frontend preview
            jd_lines = jd_text.strip().split('\n')
            return Response({"jd_lines": jd_lines})

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)